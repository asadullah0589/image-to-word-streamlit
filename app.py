import streamlit as st
import os
import requests
import base64
import json
from pathlib import Path
import tempfile
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# ==================== CONFIGURATION ====================
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
VISION_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"
# ======================================================

if not GROQ_API_KEY:
    st.error(
        "GROQ_API_KEY is not set.\n\n"
        "Add it in Streamlit Secrets (Cloud) or as an environment variable (local)."
    )
    st.stop()


st.set_page_config(
    page_title="Image to Word Converter",
    page_icon="📄",
    layout="wide"
)

st.title("📄 Image to Editable Word Converter")
st.markdown("Upload document images and convert them into **fully editable Word files**")

# ------------------------------------------------------
def encode_image_to_base64(image_path):
    with open(image_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")

def get_image_type(image_path):
    ext = Path(image_path).suffix.lower()
    return {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp"
    }.get(ext, "image/jpeg")

# ------------------------------------------------------
def extract_text_with_groq(image_path):
    base64_image = encode_image_to_base64(image_path)
    image_type = get_image_type(image_path)

    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }

    prompt = """
Analyze the document image and extract ALL text with formatting.

Return STRICT JSON in this format:

{
  "sections": [
    {
      "text": "content",
      "type": "heading1|heading2|heading3|paragraph|list_bullet|list_number|table",
      "formatting": {
        "bold": true,
        "italic": false,
        "underline": false,
        "alignment": "left|center|right|justify",
        "font_size": "large|medium|small"
      },
      "table_data": {
        "cells": [["A", "B"], ["C", "D"]]
      }
    }
  ]
}
"""

    payload = {
        "model": VISION_MODEL,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{image_type};base64,{base64_image}"
                        }
                    },
                    {
                        "type": "text",
                        "text": prompt
                    }
                ]
            }
        ],
        "temperature": 0.1,
        "max_tokens": 6000
    }

    response = requests.post(
        GROQ_API_URL,
        headers=headers,
        json=payload,
        timeout=120
    )

    if response.status_code != 200:
        st.error(response.text)
        return None

    content = response.json()["choices"][0]["message"]["content"]

    match = re.search(r"\{.*\}", content, re.DOTALL)
    if match:
        return json.loads(match.group())

    return None

# ------------------------------------------------------
def create_docx(sections, output_path):
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(12)

    for sec in sections:
        text = sec.get("text", "").strip()
        if not text:
            continue

        t = sec.get("type", "paragraph")
        f = sec.get("formatting", {})

        if t == "heading1":
            p = doc.add_heading(text, 1)
        elif t == "heading2":
            p = doc.add_heading(text, 2)
        elif t == "heading3":
            p = doc.add_heading(text, 3)
        elif t == "list_bullet":
            p = doc.add_paragraph(text, style="List Bullet")
        elif t == "list_number":
            p = doc.add_paragraph(text, style="List Number")
        elif t == "table" and sec.get("table_data"):
            cells = sec["table_data"].get("cells", [])
            if cells:
                table = doc.add_table(rows=len(cells), cols=len(cells[0]))
                for i, row in enumerate(cells):
                    for j, val in enumerate(row):
                        table.rows[i].cells[j].text = str(val)
            continue
        else:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = f.get("bold", False)
            run.italic = f.get("italic", False)
            run.underline = f.get("underline", False)

        align = f.get("alignment", "left")
        p.alignment = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }.get(align, WD_ALIGN_PARAGRAPH.LEFT)

    doc.save(output_path)

# ======================= UI ============================
uploaded_files = st.file_uploader(
    "Upload document images",
    type=["png", "jpg", "jpeg"],
    accept_multiple_files=True
)

if st.button("🚀 Convert to Word"):
    if not uploaded_files:
        st.warning("Upload at least one image")
        st.stop()

    with tempfile.TemporaryDirectory() as tmp:
        all_sections = []

        for idx, file in enumerate(uploaded_files):
            img_path = Path(tmp) / file.name
            with open(img_path, "wb") as f:
                f.write(file.getbuffer())

            st.info(f"Processing {file.name}")
            data = extract_text_with_groq(str(img_path))

            if data and "sections" in data:
                all_sections.extend(data["sections"])

        if not all_sections:
            st.error("No text extracted")
            st.stop()

        output = Path(tmp) / "converted.docx"
        create_docx(all_sections, output)

        with open(output, "rb") as f:
            st.download_button(
                "📥 Download Word File",
                f,
                file_name="converted_document.docx"
            )

st.markdown("---")
st.caption("Powered by Groq Vision + python-docx")
